from typing import Literal
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from tools.cleans import (
    extract_from_first_digit, remove_from_first_digit, starts_with_number_like,
    clean_paragraph, extract_first_text_after_number, clean_texte_special
)
from pathlib import Path
import csv

Unity = Literal["GHz", "MHz", "kHz"]
BASE_PATH = Path(__file__).resolve().parent.parent


class ExtractData:
    def __init__(self, path_file:str, unity:Unity, path_global_file:str) -> None:
        self.unity = unity
        self.path_file = path_file
        self.path_global_file = path_global_file
        self.tables = self._extract_table_in_file()
        self.paragraphs = self._extract_paragraph_in_global_file()

    def _extract_paragraph_in_global_file(self):
        """
        Retourne la liste des paragraphes du fichier global.
        """
        doc = Document(self.path_global_file)
        paragraphs = [p.text for p in doc.paragraphs]
        return paragraphs
    
    def _extract_table_in_file(self):
        """
        Retourne la liste des tableaux du fichier.
        """
    
        doc = Document(self.path_file)
        tables = [t for t in doc.tables]
        return tables
    
    def _extract_data_from_table(self, table:Table):
        """
        Retourne la liste des services dans la table.
        """
        
        lines = []
        for i_row in range(2, len(table.rows)):
            row = table.rows[i_row]
            if len(row.cells) == 5:
                cells = row.cells
                bande = f"{clean_texte_special(cells[2].text)} {self.unity}"
                services = [p.text for p in cells[3].paragraphs if p.text != " " and  not starts_with_number_like(p.text)]
                renvoie_specifique = {p: [r for r in extract_from_first_digit(p) if r] for p in services}
                text_renvoie_specifique = {rs: extract_first_text_after_number(self.paragraphs, rs) for _rs in renvoie_specifique.values() for rs in _rs}
                renvoie_global = [p for p in cells[3].paragraphs[-1].text.split(" ") if p] if starts_with_number_like(cells[3].paragraphs[-1].text) else []
                text_renvoie_global = {rg: extract_first_text_after_number(self.paragraphs, rg) for rg in renvoie_global}
                observation = cells[4].text

                for service in services:
                    greaters_lines = max(len(renvoie_specifique[service]), len(renvoie_global))
                    for i in range(greaters_lines):
                        lines.append([
                            clean_paragraph(bande),
                            clean_paragraph(remove_from_first_digit(service)),
                            clean_paragraph(renvoie_specifique[service][i]) if i < len(renvoie_specifique[service]) else "",
                            clean_paragraph(text_renvoie_specifique[renvoie_specifique[service][i]]) if i < len(renvoie_specifique[service]) else "",
                            clean_paragraph(renvoie_global[i]) if i < len(renvoie_global) else "",
                            clean_paragraph(text_renvoie_global[renvoie_global[i]]) if i < len(renvoie_global) else "",
                            clean_paragraph(observation)
                        ])

        return lines

    def extract_data(self):
        data = [
             ["bandes", "services", "renvoie_specifique", "text_renvoie_specifique", "renvoie_global", "text_renvoie_global", "observation"]
        ]
        for table in self.tables:
            data.extend(self._extract_data_from_table(table))
        return data

    def write_data_in_csv(self, path_file:str|None=None):
        path_file = path_file if path_file  else f"{BASE_PATH}/output/{self.unity}.csv"
        with open(path_file, "w", newline="") as f:
            writer = csv.writer(f)
            writer.writerows(self.extract_data())


