import csv
from pathlib import Path

from docx import Document
from docx.table import Table

from scrapping_docs.new_version.models.frequency import Unity
from scrapping_docs.new_version.tools.cleans import (
    clean_paragraph,
    clean_texte_special,
    extract_first_text_after_number,
    extract_from_first_digit,
    remove_from_first_digit,
    starts_with_number_like,
)

BASE_PATH = Path(__file__).resolve().parent.parent  # scrapping_docs/new_version/


class ExtractData:
    """
    Lit un tableau TANARES par unité (kHz/MHz/GHz) et le document global
    (qui porte les paragraphes de renvois numérotés, ex. '5.54') pour
    produire les lignes plates (bande, service, renvois) exploitées
    par BandRepository.
    """

    def __init__(self, path_file: str, unity: Unity, path_global_file: str) -> None:
        self.unity = unity
        self.path_file = path_file
        self.path_global_file = path_global_file
        self.tables = self._extract_table_in_file()
        self.paragraphs = self._extract_paragraph_in_global_file()

    def _extract_paragraph_in_global_file(self) -> list[str]:
        doc = Document(self.path_global_file)
        return [p.text for p in doc.paragraphs]

    def _extract_table_in_file(self) -> list[Table]:
        doc = Document(self.path_file)
        return list(doc.tables)

    def _extract_data_from_table(self, table: Table) -> list[list[str]]:
        lines = []
        for row in table.rows:
            cells = row.cells
            bande = f"{clean_texte_special(cells[2].text)}-{self.unity}"
            services = [
                p.text
                for p in cells[3].paragraphs
                if p.text != " "
                and not starts_with_number_like(p.text)
                and p.text != "SERVICES"
                and p.text != "COTE D’IVOIRE"
            ]
            renvoie_specifique = {
                p: [r for r in extract_from_first_digit(p) if r] for p in services
            }
            text_renvoie_specifique = {
                rs: extract_first_text_after_number(self.paragraphs, rs)
                for _rs in renvoie_specifique.values()
                for rs in _rs
            }
            renvoie_global = (
                [p for p in cells[3].paragraphs[-1].text.split(" ") if p]
                if starts_with_number_like(cells[3].paragraphs[-1].text)
                else []
            )
            text_renvoie_global = {
                rg: extract_first_text_after_number(self.paragraphs, rg)
                for rg in renvoie_global
            }
            observation = cells[4].text

            for service in services:
                nbr_lines = max(1, len(renvoie_specifique[service]), len(renvoie_global))
                for i in range(nbr_lines):
                    lines.append(
                        [
                            self.unity,
                            clean_paragraph(bande),
                            clean_paragraph(remove_from_first_digit(service)),
                            clean_paragraph(renvoie_specifique[service][i])
                            if i < len(renvoie_specifique[service])
                            else "",
                            clean_paragraph(
                                text_renvoie_specifique[renvoie_specifique[service][i]]
                            )
                            if i < len(renvoie_specifique[service])
                            else "",
                            clean_paragraph(renvoie_global[i])
                            if i < len(renvoie_global)
                            else "",
                            clean_paragraph(text_renvoie_global[renvoie_global[i]])
                            if i < len(renvoie_global)
                            else "",
                            clean_paragraph(observation),
                        ]
                    )
        return lines

    def extract_data(self) -> list[list[str]]:
        header = [
            "unity", "bandes", "services", "renvoie_specifique",
            "text_renvoie_specifique", "renvoie_global", "text_renvoie_global",
            "observation",
        ]
        data = [header]
        for table in self.tables:
            data.extend(self._extract_data_from_table(table))
        return data

    def write_data_in_csv(self, path_file: str | None = None) -> str:
        path_file = path_file or str(BASE_PATH / "output" / f"{self.unity}.csv")
        Path(path_file).parent.mkdir(parents=True, exist_ok=True)
        with open(path_file, "w", newline="") as f:
            writer = csv.writer(f)
            writer.writerows(self.extract_data())
        return path_file
