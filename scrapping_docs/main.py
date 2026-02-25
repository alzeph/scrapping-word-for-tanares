from docx import Document
from tabulate import tabulate
from tools.cleans import (
    extract_from_first_digit, remove_from_first_digit, starts_with_number_like,
    filter_exact_number_prefix, starts_with_exact_number, extract_first_text_after_number
)
import csv
doc = Document("/Users/admin/cedric/scrapping-file-word/scrapping_docs/asserts/TANARES.docx")

# for i, table in enumerate(doc.tables):
#     print(f"Table {i}")
#     for row in table.rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 print(paragraph.text)

# table = doc.tables[0]
# print(" Number of columns: ", len(table.columns))


# for row in table.rows:
#     for cell in row.cells:
#         for paragraph in cell.paragraphs:
#             print(paragraph.text)

# row = table.rows[3]
# for cell in row.cells:
#     for paragraph in cell.paragraphs:
#         print(paragraph.text)



col_tab = ["bandes", "services", "renvoie_specifique", "text_renvoie_specifique", "renvoie_global", "text_renvoie_global", "observation"]
tab = [
    col_tab,
]
print(len(doc.tables))
paragraphs = [p.text for p in doc.paragraphs]

for i_table, table in enumerate(doc.tables):
    for i_row in range(2, len(table.rows)):
        row = table.rows[i_row]
        print("cells == ",len(row.cells))
        if len(row.cells) == 5:
            cells = row.cells
            bande = cells[2].text
            services = [p.text for p in cells[3].paragraphs if p.text != " " and  not starts_with_number_like(p.text)]
            renvoie_specifique = {p: [r for r in extract_from_first_digit(p) if r] for p in services}
            text_renvoie_specifique = {rs: extract_first_text_after_number(paragraphs, rs) for _rs in renvoie_specifique.values() for rs in _rs}
            renvoie_global = [p for p in cells[3].paragraphs[-1].text.split(" ") if p] if starts_with_number_like(cells[3].paragraphs[-1].text) else []
            text_renvoie_global = {rg: extract_first_text_after_number(paragraphs, rg) for rg in renvoie_global}
            observation = cells[4].text

            lines = []
            for service in services:
                greaters_lines = max(len(renvoie_specifique[service]), len(renvoie_global))
                print("table == ",i_table)
                print("service == ",service)
                print("renvoie specifique == ",renvoie_specifique[service])
                print("renvoie global == ",renvoie_global)
                print("greaters_lines == ",greaters_lines)
                print("\n")
                for i in range(greaters_lines):
                    lines.append([
                        bande,
                        remove_from_first_digit(service),
                        renvoie_specifique[service][i] if i < len(renvoie_specifique[service]) else "",
                        text_renvoie_specifique[renvoie_specifique[service][i]] if i < len(renvoie_specifique[service]) else "",
                        renvoie_global[i] if i < len(renvoie_global) else "",
                        text_renvoie_global[renvoie_global[i]] if i < len(renvoie_global) else "",
                        observation
                    ])

            tab.extend(lines)

with open("output.csv", "w", newline="") as f:
    writer = csv.writer(f)
    writer.writerows(tab)

